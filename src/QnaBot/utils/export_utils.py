import pandas as pd
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

def export_to_pdf(qna_list, output_path="qna_output.pdf"):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    elements = []
    for idx, qa in enumerate(qna_list, 1):
        page_info = f"(Page {qa['page']}) " if qa.get("page") else ""
        elements.append(Paragraph(f"{page_info}Q{idx}: {qa['question']}", styles['Heading3']))
        elements.append(Paragraph(f"A{idx}: {qa['answer']}", styles['Normal']))
        elements.append(Spacer(1, 12))
    doc.build(elements)
    with open(output_path, "wb") as f:
        f.write(buffer.getvalue())

def export_to_csv(qna_list, output_path="qna_output.csv"):
    df = pd.DataFrame(qna_list)
    df.to_csv(output_path, index=False, encoding='utf-8')

def export_to_word(qna_list, output_path="qna_output.docx"):
    doc = Document()
    doc.add_heading("QnA Set", level=1)
    for idx, qa in enumerate(qna_list, 1):
        page_info = f"(Page {qa['page']}) " if qa.get("page") else ""
        doc.add_paragraph(f"{page_info}Q{idx}: {qa['question']}", style="List Number")
        doc.add_paragraph(f"A{idx}: {qa['answer']}")
        doc.add_paragraph("")
    doc.save(output_path)
